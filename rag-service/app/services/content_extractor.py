import html as html_lib
import json
import logging
import re
from typing import Any
from urllib.parse import parse_qs, urlparse
from xml.etree import ElementTree as ET
from xml.etree.ElementTree import ParseError

import httpx
import trafilatura
from bs4 import BeautifulSoup
from youtube_transcript_api import YouTubeTranscriptApi

try:
    from yt_dlp import YoutubeDL
except Exception:  # noqa: BLE001
    YoutubeDL = None  # type: ignore[assignment]

URL_PATTERN = re.compile(r"https?://[^\s]+", re.IGNORECASE)
logger = logging.getLogger(__name__)

_NEWS_DOMAINS = {
    "techcrunch.com", "zdnet.com", "wired.com", "theverge.com",
    "reuters.com", "bloomberg.com", "forbes.com", "cnbc.com",
    "venturebeat.com", "arstechnica.com", "engadget.com",
    "chosun.com", "joins.com", "hani.co.kr", "khan.co.kr",
    "yonhapnewstv.co.kr", "yonhap.co.kr", "ytn.co.kr",
    "etnews.com", "zdnet.co.kr", "itworld.co.kr", "aitimes.com",
    "news.naver.com", "news.daum.net",
}
_NEWS_PATH_RE = re.compile(r"/(?:news|article|story|press|release|[0-9]{4}/[0-9]{2})/", re.IGNORECASE)

_BLOG_DOMAINS = {
    "medium.com", "substack.com", "tistory.com", "velog.io",
    "brunch.co.kr", "blog.naver.com",
    "wordpress.com", "blogspot.com", "ghost.io", "hashnode.dev",
    "dev.to", "notion.so",
}
_BLOG_PATH_RE = re.compile(r"/(blog|posts?|b|writing)/", re.IGNORECASE)


def _infer_source_type(url: str) -> str:
    parsed = urlparse(url)
    host = parsed.netloc.lower().lstrip("www.")
    path = parsed.path

    if any(host == d or host.endswith("." + d) for d in _NEWS_DOMAINS):
        return "news"
    if _NEWS_PATH_RE.search(path):
        return "news"
    if any(host == d or host.endswith("." + d) for d in _BLOG_DOMAINS):
        return "blog"
    if _BLOG_PATH_RE.search(path):
        return "blog"
    return "other"


_JUNK_SIGNALS = re.compile(
    r"(등록번호|발행인|편집인|발행소|저작권법|무단\s*전재|all\s+rights\s+reserved"
    r"|copyright\s*[©ⓒ]|개인정보\s*처리방침|이용약관|광고문의|사업자등록번호)",
    re.IGNORECASE,
)

_PAYWALL_SIGNALS = re.compile(
    r"(로그인\s*후\s*이용|로그인\s*하러\s*가기|로그인\s*하시면|유료\s*회원|구독\s*후\s*이용"
    r"|subscribe\s*to\s*(read|access|continue)|this\s*content\s*is\s*(for\s*)?subscribers?"
    r"|sign\s*in\s*to\s*(read|access|continue)|premium\s*(content|article)"
    r"|회원\s*전용|유료\s*기사|전문\s*보기.*로그인|멤버십.*전용)",
    re.IGNORECASE,
)


def is_valid_content(text: str) -> bool:
    if not text or len(text.strip()) < 100:
        return False
    printable = sum(1 for c in text if c.isprintable() or c in "\n\r\t")
    if printable / len(text) < 0.85:
        return False
    words = text.split()
    junk_hits = len(_JUNK_SIGNALS.findall(text))
    if junk_hits >= 2 and len(words) < 200:
        return False
    return True


def extract_first_url(text: str) -> str | None:
    match = URL_PATTERN.search(text or "")
    if not match:
        return None
    return match.group(0).rstrip(').,\"\'')


def _youtube_video_id(url: str) -> str | None:
    parsed = urlparse(url)
    host = parsed.netloc.lower()

    if "youtu.be" in host:
        return parsed.path.strip("/") or None

    if "youtube.com" in host:
        if parsed.path == "/watch":
            return parse_qs(parsed.query).get("v", [None])[0]
        if parsed.path.startswith("/shorts/"):
            return parsed.path.split("/shorts/")[-1].split("/")[0]
        if parsed.path.startswith("/embed/"):
            return parsed.path.split("/embed/")[-1].split("/")[0]
        if parsed.path.startswith("/live/"):
            return parsed.path.split("/live/")[-1].split("/")[0]
    return None


def _clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def _dedup_transcript(text: str) -> str:
    words = text.split()
    if len(words) < 4:
        return text
    result: list[str] = [words[0]]
    i = 1
    while i < len(words):
        deduped = False
        for size in range(min(40, len(result), len(words) - i), 2, -1):
            if result[-size:] == words[i:i + size]:
                i += size
                deduped = True
                break
        if not deduped:
            result.append(words[i])
            i += 1
    return " ".join(result)


def _extract_html_title(html: str) -> str:
    soup = BeautifulSoup(html, "html.parser")
    title = (soup.title.string or "").strip() if soup.title else ""
    return title or "Untitled"


_BROWSER_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    ),
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "ko-KR,ko;q=0.9,en-US;q=0.8,en;q=0.7",
    "Accept-Encoding": "gzip, deflate",
}


def _fetch_html(url: str, timeout: int) -> tuple[str, str]:
    with httpx.Client(
        timeout=timeout,
        follow_redirects=True,
        headers=_BROWSER_HEADERS,
    ) as client:
        response = client.get(url)

        if response.status_code == 403:
            response = client.get(
                url,
                headers={
                    **_BROWSER_HEADERS,
                    "User-Agent": "Googlebot/2.1 (+http://www.google.com/bot.html)",
                },
            )

        if response.status_code == 403:
            raise ValueError(
                f"해당 페이지는 외부 접근을 차단하고 있어 내용을 가져올 수 없습니다. (403 Forbidden)\n"
                f"URL: {response.url}"
            )

        response.raise_for_status()
        return str(response.url), response.text


class ExtractedContent:
    def __init__(self, url: str, source_type: str, title: str, content: str) -> None:
        self.url = url
        self.source_type = source_type
        self.title = title
        self.content = content


def _extract_from_web(url: str, timeout: int) -> ExtractedContent:
    final_url, html = _fetch_html(url, timeout)

    redirected_video_id = _youtube_video_id(final_url)
    if redirected_video_id:
        return _extract_from_youtube(final_url, timeout)

    # 페이월 조기 감지: raw HTML 기준으로 로그인/구독 유도 문구 반복 확인
    if len(_PAYWALL_SIGNALS.findall(html)) >= 3:
        parsed_host = urlparse(final_url).netloc.lstrip("www.")
        raise ValueError(
            f"해당 기사는 로그인 또는 유료 구독이 필요한 콘텐츠입니다.\n"
            f"출처: {parsed_host}\n\n"
            f"이 사이트는 비로그인 상태에서 전문 접근을 차단하고 있어 요약할 수 없습니다. "
            f"브라우저에서 직접 로그인 후 읽으시거나, "
            f"해당 언론사 앱/웹사이트에서 구독 후 이용해주세요."
        )

    title = _extract_html_title(html)
    content = trafilatura.extract(
        html,
        include_comments=False,
        include_tables=False,
        favor_precision=True,
    ) or ""

    cleaned = _clean_text(content)

    # trafilatura 결과가 너무 짧으면 CSS 셀렉터 기반 fallback 시도
    if len(cleaned) < 1000:
        soup = BeautifulSoup(html, "html.parser")

        def _find_article(soup: BeautifulSoup) -> str:
            # 클래스 부분 매칭으로 article 요소 탐색
            _KEYWORDS = ("tt_article_useless", "article_view", "entry-content",
                         "post-content", "post_content", "blog-post", "article-body")
            for kw in _KEYWORDS:
                el = soup.find(lambda tag: bool(
                    tag.get("class") and any(kw in c for c in tag.get("class", []))
                ))
                if el:
                    return _clean_text(el.get_text(" "))
            # id 기반 fallback
            for id_ in ("content", "article", "post", "main-content"):
                el = soup.find(id=id_)
                if el:
                    return _clean_text(el.get_text(" "))
            # <article> 태그
            el = soup.find("article")
            if el:
                return _clean_text(el.get_text(" "))
            return ""

        candidate = _find_article(soup)
        if len(candidate) > len(cleaned):
            cleaned = candidate

    if not cleaned:
        cleaned = _clean_text(BeautifulSoup(html, "html.parser").get_text(" "))

    return ExtractedContent(
        url=final_url,
        source_type=_infer_source_type(final_url),
        title=title,
        content=cleaned,
    )


def _parse_vtt(raw: str) -> str:
    lines = raw.splitlines()
    parts: list[str] = []
    for line in lines:
        s = line.strip()
        if not s:
            continue
        upper = s.upper()
        if upper.startswith("WEBVTT") or upper.startswith("NOTE"):
            continue
        if upper.startswith("KIND:") or upper.startswith("LANGUAGE:"):
            continue
        if "-->" in s:
            continue
        if s.isdigit():
            continue
        parts.append(s)

    text = " ".join(parts)
    text = re.sub(r"<[^>]+>", " ", text)
    return html_lib.unescape(text)


def _parse_xml_caption(raw: str) -> str:
    try:
        root = ET.fromstring(raw)
    except ParseError:
        return ""

    parts: list[str] = []
    for node in root.findall(".//text"):
        item = "".join(node.itertext()).strip()
        if item:
            parts.append(item)

    if not parts:
        return ""
    return html_lib.unescape(" ".join(parts))


def _parse_json3_caption(raw: str) -> str:
    try:
        data = json.loads(raw)
    except json.JSONDecodeError:
        return ""

    parts: list[str] = []
    events = data.get("events", []) if isinstance(data, dict) else []
    for ev in events:
        if not isinstance(ev, dict):
            continue
        segs = ev.get("segs", [])
        if not isinstance(segs, list):
            continue
        for seg in segs:
            if not isinstance(seg, dict):
                continue
            txt = seg.get("utf8", "")
            if isinstance(txt, str) and txt:
                parts.append(txt.replace("\n", " "))

    if not parts:
        return ""
    return html_lib.unescape(" ".join(parts))


def _decode_subtitle(raw: str, ext: str) -> str:
    ext_norm = (ext or "").lower()

    if ext_norm == "json3":
        parsed = _parse_json3_caption(raw)
        if parsed:
            return parsed

    if ext_norm in {"srv3", "ttml", "xml"}:
        parsed = _parse_xml_caption(raw)
        if parsed:
            return parsed

    stripped = raw.lstrip()
    if stripped.startswith("{"):
        parsed = _parse_json3_caption(raw)
        if parsed:
            return parsed

    if stripped.startswith("<"):
        parsed = _parse_xml_caption(raw)
        if parsed:
            return parsed

    return _parse_vtt(raw)


def _build_yt_dlp_candidates(info: dict[str, Any]) -> list[dict[str, str]]:
    subtitles = info.get("subtitles") if isinstance(info, dict) else None
    auto = info.get("automatic_captions") if isinstance(info, dict) else None

    source_buckets: list[tuple[str, dict[str, Any]]] = []
    if isinstance(subtitles, dict):
        source_buckets.append(("subtitles", subtitles))
    if isinstance(auto, dict):
        source_buckets.append(("automatic_captions", auto))

    def ordered_lang_keys(bucket: dict[str, Any]) -> list[str]:
        preferred = ["ko", "ko-KR", "en", "en-US"]
        keys = list(bucket.keys())
        out: list[str] = []
        for p in preferred:
            if p in bucket and p not in out:
                out.append(p)
        for k in keys:
            if (k.startswith("ko") or k.startswith("en")) and k not in out:
                out.append(k)
        for k in keys:
            if k not in out:
                out.append(k)
        return out

    ext_rank = {"vtt": 0, "srv3": 1, "json3": 2, "ttml": 3, "xml": 4}
    scored: list[tuple[int, int, int, dict[str, str]]] = []

    for source_idx, (source_name, bucket) in enumerate(source_buckets):
        for lang_idx, lang in enumerate(ordered_lang_keys(bucket)):
            tracks = bucket.get(lang)
            if not isinstance(tracks, list):
                continue
            for track in tracks:
                if not isinstance(track, dict):
                    continue
                track_url = track.get("url")
                if not isinstance(track_url, str) or not track_url:
                    continue
                ext = str(track.get("ext", "")).lower()
                payload = {
                    "source": source_name,
                    "lang": lang,
                    "ext": ext,
                    "url": track_url,
                }
                scored.append((source_idx, lang_idx, ext_rank.get(ext, 9), payload))

    scored.sort(key=lambda x: (x[0], x[1], x[2]))
    return [item[3] for item in scored]


def _extract_with_yt_dlp(url: str, video_id: str, timeout: int) -> str:
    if YoutubeDL is None:
        logger.warning("YouTube fallback unavailable: video_id=%s method=yt_dlp reason=module_missing", video_id)
        return ""

    try:
        with YoutubeDL({"skip_download": True, "quiet": True, "no_warnings": True, "extract_flat": False}) as ydl:
            info = ydl.extract_info(url, download=False)
    except Exception as exc:  # noqa: BLE001
        logger.warning("YouTube fallback failed: video_id=%s method=yt_dlp reason=%s", video_id, exc.__class__.__name__)
        return ""

    if not isinstance(info, dict):
        return ""

    candidates = _build_yt_dlp_candidates(info)
    with httpx.Client(timeout=timeout, follow_redirects=True) as client:
        for idx, candidate in enumerate(candidates, start=1):
            try:
                response = client.get(candidate["url"])
                response.raise_for_status()
                parsed = _decode_subtitle(response.text, candidate["ext"])
                cleaned = _clean_text(parsed)
                if cleaned:
                    logger.info(
                        "YouTube extraction success: video_id=%s method=yt_dlp lang=%s track_index=%d chars=%d",
                        video_id, candidate["lang"], idx, len(cleaned),
                    )
                    return cleaned
            except Exception as exc:  # noqa: BLE001
                logger.warning("YouTube yt-dlp track failed: video_id=%s track_index=%d reason=%s", video_id, idx, exc.__class__.__name__)

    return ""


def _get_youtube_title(video_id: str, url: str, timeout: int) -> str:
    try:
        resp = httpx.get(
            f"https://www.youtube.com/oembed?url={url}&format=json",
            timeout=timeout,
            follow_redirects=True,
        )
        resp.raise_for_status()
        title = resp.json().get("title", "").strip()
        if title:
            return title
    except Exception:  # noqa: BLE001
        pass
    return f"YouTube Video ({video_id})"


def _extract_from_youtube(url: str, timeout: int) -> ExtractedContent:
    video_id = _youtube_video_id(url)
    if not video_id:
        raise ValueError("유효한 유튜브 영상 URL이 아닙니다.")

    title = _get_youtube_title(video_id, url, timeout)
    logger.info("YouTube extraction start: video_id=%s url=%s title=%r", video_id, url, title)

    cleaned = ""
    _api = YouTubeTranscriptApi()

    def _join_transcript(items: Any) -> str:
        parts = []
        for item in items:
            text = item.get("text", "") if isinstance(item, dict) else getattr(item, "text", "")
            if text:
                parts.append(str(text))
        return " ".join(parts)

    # 1차: ko/en 자막 직접 요청
    try:
        fetched = _api.fetch(video_id, languages=["ko", "en"])
        joined = _join_transcript(fetched)
        cleaned = _dedup_transcript(_clean_text(joined))
        logger.info("YouTube extraction success: video_id=%s method=transcript_api chars=%d", video_id, len(cleaned))
    except ParseError:
        logger.warning("YouTube extraction failed: video_id=%s method=transcript_api reason=parse_error", video_id)
    except Exception as exc:  # noqa: BLE001
        logger.warning("YouTube extraction failed: video_id=%s method=transcript_api reason=%s", video_id, exc.__class__.__name__)

    if cleaned:
        return ExtractedContent(url=url, source_type="youtube", title=title, content=cleaned)

    logger.info("YouTube extraction fallback start: video_id=%s method=yt_dlp", video_id)
    cleaned = _dedup_transcript(_extract_with_yt_dlp(url, video_id, timeout))
    if cleaned:
        return ExtractedContent(url=url, source_type="youtube", title=title, content=cleaned)

    raise ValueError(
        "유튜브 자막 텍스트를 추출하지 못해 요약할 수 없습니다. "
        "비공개·연령 제한·자막 비활성화 영상일 수 있습니다."
    )


_FILE_CONTENT_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "docx",
}


def _detect_file_type(url: str, timeout: int) -> str | None:
    """URL이 PDF/docx 파일인지 확인. 확장자 우선, 없으면 HEAD 요청으로 Content-Type 확인."""
    # URL 경로에서 확장자 추출 (쿼리스트링 제거)
    path = urlparse(url).path.lower()
    ext = path.rsplit(".", 1)[-1] if "." in path else ""
    if ext == "pdf":
        return "pdf"
    if ext in ("docx", "doc"):
        return "docx"

    # 확장자로 판단 불가 → HEAD 요청으로 Content-Type 확인
    try:
        with httpx.Client(timeout=10, follow_redirects=True) as client:
            resp = client.head(url, headers=_BROWSER_HEADERS)
            ct = resp.headers.get("content-type", "").split(";")[0].strip().lower()
            return _FILE_CONTENT_TYPES.get(ct)
    except Exception:  # noqa: BLE001
        return None


def _fetch_file_bytes(url: str, timeout: int) -> bytes:
    with httpx.Client(timeout=timeout, follow_redirects=True, headers=_BROWSER_HEADERS) as client:
        resp = client.get(url)
        resp.raise_for_status()
        return resp.content


def _fetch_gdrive_bytes(file_id: str, timeout: int) -> tuple[bytes, str]:
    """Google Drive 파일을 다운로드. 바이러스 스캔 확인 페이지 처리 포함.
    Returns: (file_bytes, content_type)
    """
    download_url = f"https://drive.google.com/uc?export=download&id={file_id}"
    with httpx.Client(timeout=timeout, follow_redirects=True, headers=_BROWSER_HEADERS) as client:
        resp = client.get(download_url)
        resp.raise_for_status()
        ct = resp.headers.get("content-type", "").split(";")[0].strip().lower()

        # 대용량 파일의 경우 바이러스 스캔 확인 HTML 페이지가 반환됨
        if "text/html" in ct:
            logger.info("Google Drive confirmation page detected: file_id=%s", file_id)
            # confirm 토큰 추출 시도 (여러 패턴)
            confirm_token = None
            m = re.search(r'name="confirm"\s+value="([^"]+)"', resp.text)
            if m:
                confirm_token = m.group(1)
            if not confirm_token:
                m = re.search(r'[?&]confirm=([^&"&\s]+)', resp.text)
                if m:
                    confirm_token = m.group(1)
            if not confirm_token:
                # 새 방식: /uc?id=... 링크에서 confirm 추출
                m = re.search(r'href="(/uc\?export=download[^"]*confirm=[^"]+)"', resp.text)
                if m:
                    confirm_path = m.group(1).replace("&amp;", "&")
                    confirm_url = f"https://drive.google.com{confirm_path}"
                    resp = client.get(confirm_url)
                    resp.raise_for_status()
                    ct = resp.headers.get("content-type", "").split(";")[0].strip().lower()
                    return resp.content, ct

            if confirm_token:
                confirm_url = f"https://drive.google.com/uc?export=download&id={file_id}&confirm={confirm_token}"
                resp = client.get(confirm_url)
                resp.raise_for_status()
                ct = resp.headers.get("content-type", "").split(";")[0].strip().lower()
            else:
                raise ValueError(
                    "Google Drive 파일에 접근할 수 없습니다. "
                    "파일이 '링크가 있는 모든 사용자' 공개 상태인지 확인하세요."
                )

        return resp.content, ct


def _filename_from_url(url: str, fallback_ext: str) -> str:
    path = urlparse(url).path
    name = path.rsplit("/", 1)[-1].split("?")[0].strip()
    return name if name else f"document.{fallback_ext}"


_GDOCS_RE = re.compile(
    r"https://docs\.google\.com/document/d/([A-Za-z0-9_-]+)",
    re.IGNORECASE,
)
_GSHEETS_RE = re.compile(
    r"https://docs\.google\.com/spreadsheets/d/([A-Za-z0-9_-]+)",
    re.IGNORECASE,
)
_GDRIVE_RE = re.compile(
    r"https://drive\.google\.com/(?:file/d/|open\?id=)([A-Za-z0-9_-]+)",
    re.IGNORECASE,
)


def _gdocs_export_url(url: str) -> tuple[str, str] | None:
    """Google Docs/Sheets URL을 export URL과 포맷으로 변환."""
    m = _GDOCS_RE.search(url)
    if m:
        doc_id = m.group(1)
        return (f"https://docs.google.com/document/d/{doc_id}/export?format=docx", "docx")
    m = _GSHEETS_RE.search(url)
    if m:
        doc_id = m.group(1)
        return (f"https://docs.google.com/spreadsheets/d/{doc_id}/export?format=xlsx", "xlsx")
    return None


def _gdrive_file_id(url: str) -> str | None:
    """Google Drive 공유 URL에서 file ID 추출."""
    m = _GDRIVE_RE.search(url)
    return m.group(1) if m else None


def extract_content(url: str, timeout: int) -> ExtractedContent:
    if _youtube_video_id(url):
        logger.info("Extractor selected: youtube_direct url=%s", url)
        return _extract_from_youtube(url, timeout)

    # Google Docs / Sheets → export로 변환
    gdocs = _gdocs_export_url(url)
    if gdocs:
        export_url, fmt = gdocs
        logger.info("Extractor selected: google_docs url=%s export_fmt=%s", url, fmt)
        file_bytes = _fetch_file_bytes(export_url, timeout)
        filename = f"google_doc.{fmt}"
        if fmt == "docx":
            return extract_from_docx(file_bytes, filename)
        raise ValueError("Google Sheets(xlsx) 요약은 아직 지원하지 않습니다.")

    # Google Drive 파일 (PDF/docx 등)
    drive_id = _gdrive_file_id(url)
    if drive_id:
        logger.info("Extractor selected: google_drive url=%s file_id=%s", url, drive_id)
        file_bytes, ct = _fetch_gdrive_bytes(drive_id, timeout)
        logger.info("Google Drive download done: file_id=%s content_type=%s bytes=%d", drive_id, ct, len(file_bytes))
        if ct == "application/pdf":
            return extract_from_pdf(file_bytes, f"gdrive_{drive_id}.pdf")
        if ct in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"):
            return extract_from_docx(file_bytes, f"gdrive_{drive_id}.docx")
        # Content-Type 불명확 시 PDF 먼저 시도
        try:
            return extract_from_pdf(file_bytes, f"gdrive_{drive_id}.pdf")
        except Exception:  # noqa: BLE001
            return extract_from_docx(file_bytes, f"gdrive_{drive_id}.docx")

    file_type = _detect_file_type(url, timeout)
    if file_type == "pdf":
        filename = _filename_from_url(url, "pdf")
        logger.info("Extractor selected: pdf url=%s filename=%s", url, filename)
        file_bytes = _fetch_file_bytes(url, timeout)
        return extract_from_pdf(file_bytes, filename)
    if file_type == "docx":
        filename = _filename_from_url(url, "docx")
        logger.info("Extractor selected: docx url=%s filename=%s", url, filename)
        file_bytes = _fetch_file_bytes(url, timeout)
        return extract_from_docx(file_bytes, filename)

    logger.info("Extractor selected: web_or_redirect url=%s", url)
    return _extract_from_web(url, timeout)


def extract_from_pdf(file_bytes: bytes, filename: str) -> ExtractedContent:
    """PDF 파일 바이트에서 텍스트를 추출합니다. PyMuPDF 우선, 실패 시 pypdf 폴백."""
    title = filename.rsplit(".", 1)[0]

    # PyMuPDF (fitz) — 한글 폰트 포함 PDF에 더 강함
    try:
        import fitz  # pymupdf
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        parts: list[str] = []
        for page in doc:
            text = page.get_text("text")
            if text and text.strip():
                parts.append(text.strip())
        doc.close()
        content = _clean_text("\n".join(parts))
        if content:
            logger.info("PDF extraction done (fitz): filename=%s pages=%d chars=%d preview=%r", filename, len(parts), len(content), content[:500])
            return ExtractedContent(url=f"upload://{filename}", source_type="pdf", title=title, content=content)
        logger.warning("PDF extraction (fitz) returned empty, falling back to pypdf: filename=%s", filename)
    except Exception as exc:  # noqa: BLE001
        logger.warning("PDF extraction (fitz) failed: filename=%s reason=%s", filename, exc)

    # pypdf 폴백
    from io import BytesIO
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("pypdf 또는 pymupdf 패키지가 설치되지 않았습니다.") from exc

    reader = PdfReader(BytesIO(file_bytes))
    parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text and text.strip():
            parts.append(text.strip())

    content = _clean_text("\n".join(parts))
    logger.info("PDF extraction done (pypdf): filename=%s pages=%d chars=%d", filename, len(reader.pages), len(content))
    return ExtractedContent(url=f"upload://{filename}", source_type="pdf", title=title, content=content)


def extract_from_excel(file_bytes: bytes, filename: str) -> ExtractedContent:
    """Excel(.xlsx) 파일 바이트에서 텍스트를 추출합니다."""
    from io import BytesIO

    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("openpyxl 패키지가 설치되지 않았습니다.") from exc

    wb = load_workbook(BytesIO(file_bytes), read_only=True, data_only=True)
    parts: list[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        sheet_rows: list[str] = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                sheet_rows.append("\t".join(cells))
        if sheet_rows:
            parts.append(f"[Sheet: {sheet_name}]\n" + "\n".join(sheet_rows))

    wb.close()
    title = filename.rsplit(".", 1)[0]
    content = _clean_text("\n\n".join(parts))
    logger.info("Excel extraction done: filename=%s sheets=%d chars=%d", filename, len(wb.sheetnames), len(content))
    return ExtractedContent(url=f"upload://{filename}", source_type="xlsx", title=title, content=content)


def extract_from_docx(file_bytes: bytes, filename: str) -> ExtractedContent:
    """Word(.docx) 파일 바이트에서 텍스트를 추출합니다."""
    from io import BytesIO

    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx 패키지가 설치되지 않았습니다.") from exc

    doc = Document(BytesIO(file_bytes))
    parts: list[str] = []

    # 본문 단락
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # 표 안 텍스트도 포함
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in parts:
                    parts.append(text)

    title = filename.rsplit(".", 1)[0]
    content = _clean_text("\n".join(parts))
    logger.info("DOCX extraction done: filename=%s paragraphs=%d chars=%d", filename, len(doc.paragraphs), len(content))
    return ExtractedContent(url=f"upload://{filename}", source_type="docx", title=title, content=content)
